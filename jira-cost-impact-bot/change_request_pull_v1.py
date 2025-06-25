import os
import re
import datetime
from collections import defaultdict
from slack_sdk import WebClient
from slack_sdk.errors import SlackApiError
import requests
from docx import Document

# ========== CONFIGURATION ==========
SLACK_BOT_TOKEN = "YOUR-SLACK-BOT-TOKEN"
CHANNEL_ID = "YOUR-CHANNEL-ID"
JIRA_DOMAIN = "your-domain.atlassian.net"
JIRA_EMAIL = "your-email@domain.com"
JIRA_API_TOKEN = "YOUR-JIRA-API-TOKEN"

CUSTOM_COST_FIELD_ID = "customfield_10285"
CUSTOM_CLOUD_REGIONS_ID = "customfield_10267"
CUSTOM_CHANGE_START_ID = "customfield_10266"
CUSTOM_CHANGE_FINISH_ID = "customfield_10272"

DEPLOYMENT_SUCCESS_STATUSES = ["Deployment Success"]

OUTPUT_FILENAME = "Jira_Cost_Report.docx"
# ===================================

slack_client = WebClient(token=SLACK_BOT_TOKEN)

def get_messages_from_past_week():
    now = datetime.datetime.now()
    one_week_ago = now - datetime.timedelta(days=7)
    oldest_ts = one_week_ago.timestamp()
    latest_ts = now.timestamp()

    messages = []
    has_more = True
    next_cursor = None

    print(f"📆 Fetching messages from: {one_week_ago.strftime('%Y-%m-%d')} to {now.strftime('%Y-%m-%d')}")

    while has_more:
        try:
            response = slack_client.conversations_history(
                channel=CHANNEL_ID,
                oldest=str(oldest_ts),
                latest=str(latest_ts),
                limit=200,
                cursor=next_cursor
            )
            messages.extend(response["messages"])
            has_more = response.get("has_more", False)
            next_cursor = response.get("response_metadata", {}).get("next_cursor")
        except SlackApiError as e:
            print(f"❌ Slack API error: {e.response['error']}")
            break

    return messages

def extract_jira_keys(messages):
    jira_keys = set()
    pattern = re.compile(r"\b[A-Z][A-Z0-9]+-\d+\b")

    for msg in messages:
        found = pattern.findall(msg.get("text", ""))
        jira_keys.update(found)

    return list(jira_keys)

def fetch_jira_info(issue_key):
    url = f"https://{JIRA_DOMAIN}/rest/api/3/issue/{issue_key}"
    auth = (JIRA_EMAIL, JIRA_API_TOKEN)
    headers = {"Accept": "application/json"}

    response = requests.get(url, headers=headers, auth=auth)
    if response.status_code != 200:
        print(f"⚠️ Error fetching {issue_key}: {response.status_code}")
        return None

    data = response.json()
    fields = data.get("fields", {})

    # Extract and format cloud regions
    cloud_regions_raw = fields.get(CUSTOM_CLOUD_REGIONS_ID, [])
    cloud_regions = ", ".join([r.get("value") for r in cloud_regions_raw]) if isinstance(cloud_regions_raw, list) else str(cloud_regions_raw)

    return {
        "key": issue_key,
        "summary": fields.get("summary", "No summary"),
        "cost": fields.get(CUSTOM_COST_FIELD_ID, "Not Set"),
        "status": fields.get("status", {}).get("name", "Unknown"),
        "cloud_regions": cloud_regions,
        "start_time": fields.get(CUSTOM_CHANGE_START_ID, "Not Set"),
        "finish_time": fields.get(CUSTOM_CHANGE_FINISH_ID, "Not Set"),
    }

def write_to_docx(grouped_data, filename=OUTPUT_FILENAME):
    doc = Document()
    doc.add_heading("Weekly Jira Cost Impact Report", level=1)

    # Sort statuses to show deployment successes first
    for status in sorted(grouped_data.keys(), key=lambda s: (s not in DEPLOYMENT_SUCCESS_STATUSES, s)):
        doc.add_heading(f"Status: {status}", level=2)

        for item in grouped_data[status]:
            doc.add_paragraph(f"Issue: {item['key']}")
            doc.add_paragraph(f"Summary: {item['summary']}")
            doc.add_paragraph(f"Cost Impact Estimation ($): {item['cost']}")
            doc.add_paragraph(f"Cloud Regions: {item['cloud_regions']}")
            doc.add_paragraph(f"Change Start Time: {item['start_time']}")
            doc.add_paragraph(f"Change Estimated Finish Time: {item['finish_time']}")
            doc.add_paragraph("-" * 40)

    doc.save(filename)
    print(f"\n✅ Report saved as: {filename}")

def main():
    print("📥 Fetching Slack messages from past 7 days...")
    messages = get_messages_from_past_week()

    print(f"🔍 Extracting Jira keys from {len(messages)} messages...")
    jira_keys = extract_jira_keys(messages)
    print(f"🎯 Found {len(jira_keys)} unique Jira tickets.")

    grouped_data = defaultdict(list)

    for key in jira_keys:
        print(f"🔄 Fetching data for {key}...")
        info = fetch_jira_info(key)
        if info:
            grouped_data[info["status"]].append(info)

    print("📝 Writing results to Word document grouped by Status...")
    write_to_docx(grouped_data)

if __name__ == "__main__":
    main()
